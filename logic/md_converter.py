# ==================================================================
# File: logic/md_converter.py
# Description: Markdown Converter - Worker thread for Word/Excel conversion
# ==================================================================

import os
import tempfile
from pathlib import Path

from PyQt6.QtCore import QThread, pyqtSignal

from core.logger import Logger


class ConversionWorker(QThread):
    """Worker thread for markdown conversion to keep UI responsive"""
    
    @staticmethod
    def fix_markdown_tables(content: str) -> str:
        """Ensure tables have a preceding blank line so strict parsers recognize them."""
        lines = content.split('\n')
        fixed_lines = []
        in_code_block = False
        
        for i, line in enumerate(lines):
            if line.strip().startswith('```'):
                in_code_block = not in_code_block
                
            if not in_code_block and line.strip().startswith('|'):
                if i > 0:
                    prev_line = lines[i-1].strip()
                    if prev_line and not prev_line.startswith('|') and not prev_line.startswith('```'):
                        fixed_lines.append('')
            
            fixed_lines.append(line)
            
        return '\n'.join(fixed_lines)

    
    progress = pyqtSignal(int)
    status = pyqtSignal(str)
    finished = pyqtSignal(bool, str)
    
    def __init__(self, input_file: str, output_file: str, conversion_type: str, use_highlighting: bool = True, paper_size: str = "A4", orientation: str = "Portrait", margin: str = "Normal", custom_margins: dict = None, excel_sheet_mode: str = "📊 One table per sheet"):
        super().__init__()
        self.input_file = input_file
        self.output_file = output_file
        self.conversion_type = conversion_type
        self.use_highlighting = use_highlighting
        self.paper_size = paper_size
        self.orientation = orientation
        self.margin = margin
        self.custom_margins = custom_margins if custom_margins else {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54}
        self.excel_sheet_mode = excel_sheet_mode
        self.logger = Logger()
    
    def run(self):
        """Run conversion in background thread"""
        try:
            self.status.emit("Loading file...")
            self.progress.emit(20)
            
            if self.conversion_type == "Word":
                self.convert_to_word()
            elif self.conversion_type == "Excel":
                self.convert_to_excel()
            elif self.conversion_type == "PDF":
                self.convert_to_pdf()
            elif self.conversion_type == "DocxToMd":
                self.convert_docx_to_md()
            elif self.conversion_type == "XlsxToMd":
                self.convert_xlsx_to_md()
            else:
                self.finished.emit(False, f"Unknown conversion type: {self.conversion_type}")
                
        except Exception as e:
            self.logger.error(f"Conversion error: {str(e)}")
            self.finished.emit(False, str(e))
    
    def convert_to_pdf(self):
        """Convert markdown to PDF using Pandoc (HTML) + WeasyPrint"""
        try:
            import pypandoc
            from weasyprint import HTML, CSS
            
            self.logger.info(f"Starting PDF conversion: {self.input_file}")
            self.status.emit("Converting Markdown to HTML...")
            self.progress.emit(30)
            
            extra_args = ['--standalone']
            if not self.use_highlighting:
                extra_args.append('--no-highlight')
                
            # Read the markdown file
            with open(self.input_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Write to a temporary file
            temp_md = Path(tempfile.gettempdir()) / f"temp_{Path(self.input_file).stem}.md"
            with open(temp_md, 'w', encoding='utf-8') as f:
                f.write(self.fix_markdown_tables(content))
            
            try:
                html_content = pypandoc.convert_file(str(temp_md), 'html', format='gfm', extra_args=extra_args)
            finally:
                # Clean up temp file
                if os.path.exists(temp_md):
                    os.unlink(str(temp_md))
                    
            self.status.emit("Rendering PDF with WeasyPrint...")
            self.progress.emit(60)
            
            # Map margins
            if self.margin == "Custom":
                unit = "in" if self.custom_margins.get("is_inch", False) else "cm"
                top = f"{self.custom_margins.get('top', 2.54)}{unit}"
                right = f"{self.custom_margins.get('right', 2.54)}{unit}"
                bottom = f"{self.custom_margins.get('bottom', 2.54)}{unit}"
                left = f"{self.custom_margins.get('left', 2.54)}{unit}"
            else:
                margin_val = "1in"
                if self.margin == "Narrow": margin_val = "0.5in"
                elif self.margin == "Wide": margin_val = "2in"
                top = right = bottom = left = margin_val
            
            orientation_str = "landscape" if self.orientation == "Landscape" else "portrait"
            
            # Inject CSS for page formatting and table rules
            css_fix = f"""<style>
            @page {{
                size: {self.paper_size} {orientation_str};
                margin: {top} {right} {bottom} {left};
            }}
            body {{ margin-top: 0 !important; padding-top: 0 !important; font-family: sans-serif; }}
            table {{ 
                width: 100%; 
                max-width: 100%; 
                table-layout: auto; 
                font-size: 8pt; 
                border-collapse: collapse;
                page-break-inside: auto;
            }}
            tr {{ 
                page-break-inside: avoid; 
                page-break-after: auto; 
            }}
            th, td {{ 
                word-break: break-word; 
                overflow-wrap: break-word; 
                white-space: normal; 
                border: 1px solid #ddd;
                padding: 4px;
            }}
            </style></head>"""
            html_content = html_content.replace("</head>", css_fix)
            
            HTML(string=html_content).write_pdf(self.output_file)
                
            self.progress.emit(100)
            status = "with syntax highlighting" if self.use_highlighting else ""
            self.logger.info(f"PDF conversion successful: {self.output_file}")
            self.finished.emit(True, f"Successfully converted to PDF {status}:\n{self.output_file}")
            
        except ImportError:
            self.logger.error("WeasyPrint is not installed or missing GTK3 dependencies")
            self.finished.emit(False, "PDF conversion failed: WeasyPrint is not installed or missing GTK3 dependencies.\n\nPlease install 'weasyprint' and ensure GTK3 is installed on Windows.")
        except Exception as e:
            self.logger.error(f"PDF conversion failed: {str(e)}")
            self.finished.emit(False, f"PDF conversion failed: {str(e)}")

    def convert_to_word(self):
        """Convert markdown to Word using pypandoc_binary"""
        try:
            import pypandoc
            
            self.logger.info(f"Starting Word conversion: {self.input_file}")
            self.status.emit("Converting to Word using Pandoc...")
            self.progress.emit(30)
            
            self.status.emit("Running Pandoc conversion...")
            self.progress.emit(50)
            
            # Read the markdown file
            with open(self.input_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Write to a temporary file
            temp_md = Path(tempfile.gettempdir()) / f"temp_{Path(self.input_file).stem}.md"
            with open(temp_md, 'w', encoding='utf-8') as f:
                f.write(self.fix_markdown_tables(content))
            
            try:
                extra_args = ['--standalone']
                if not self.use_highlighting:
                    extra_args.append('--no-highlight')
                    
                # Convert the temporary file
                pypandoc.convert_file(
                    str(temp_md),
                    'docx',
                    format='gfm',
                    outputfile=self.output_file,
                    extra_args=extra_args
                )
            finally:
                # Clean up temp file
                if os.path.exists(temp_md):
                    os.unlink(str(temp_md))
            
            self.status.emit("Applying page formatting...")
            self.progress.emit(80)
            
            try:
                import docx
                from docx.shared import Inches, Mm
                from docx.enum.section import WD_ORIENT
                
                doc = docx.Document(self.output_file)
                for section in doc.sections:
                    # Apply Paper Size
                    if self.paper_size == "A4":
                        page_width = Mm(210)
                        page_height = Mm(297)
                    elif self.paper_size == "Legal":
                        page_width = Inches(8.5)
                        page_height = Inches(14)
                    else: # Letter
                        page_width = Inches(8.5)
                        page_height = Inches(11)
                    
                    # Apply Orientation
                    if self.orientation == "Landscape":
                        section.orientation = WD_ORIENT.LANDSCAPE
                        section.page_width = page_height
                        section.page_height = page_width
                    else:
                        section.orientation = WD_ORIENT.PORTRAIT
                        section.page_width = page_width
                        section.page_height = page_height
                    
                    # Apply Margins
                    if self.margin == "Narrow":
                        m = Inches(0.5)
                        section.top_margin = m
                        section.bottom_margin = m
                        section.left_margin = m
                        section.right_margin = m
                    elif self.margin == "Wide":
                        section.top_margin = Inches(1)
                        section.bottom_margin = Inches(1)
                        section.left_margin = Inches(2)
                        section.right_margin = Inches(2)
                    elif self.margin == "Custom":
                        from docx.shared import Cm, Inches
                        is_inch = self.custom_margins.get("is_inch", False)
                        unit_func = Inches if is_inch else Cm
                        section.top_margin = unit_func(self.custom_margins.get("top", 2.54))
                        section.bottom_margin = unit_func(self.custom_margins.get("bottom", 2.54))
                        section.left_margin = unit_func(self.custom_margins.get("left", 2.54))
                        section.right_margin = unit_func(self.custom_margins.get("right", 2.54))
                    else: # Normal
                        section.top_margin = Inches(1)
                        section.bottom_margin = Inches(1)
                        section.left_margin = Inches(1)
                        section.right_margin = Inches(1)
                        
                # Fix tables being cut off: Auto resize and word wrap
                from docx.enum.table import WD_TABLE_ALIGNMENT
                from docx.oxml.shared import OxmlElement
                from docx.oxml.ns import qn
                
                for table in doc.tables:
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.autofit = False
                    table.allow_autofit = False
                    
                    # Explicitly calculate column widths so it never cuts off
                    page_width = section.page_width - section.left_margin - section.right_margin
                    if table.columns and len(table.columns) > 0:
                        col_width = int(page_width / len(table.columns))
                        
                        tblGrid = table._tbl.find(qn('w:tblGrid'))
                        if tblGrid is not None:
                            for gridCol in tblGrid.findall(qn('w:gridCol')):
                                gridCol.set(qn('w:w'), str(col_width))
                        
                        from docx.shared import Pt
                        for row in table.rows:
                            for cell in row.cells:
                                cell.width = col_width
                                tcPr = cell._tc.get_or_add_tcPr()
                                tcW = tcPr.find(qn('w:tcW'))
                                if tcW is not None:
                                    tcW.set(qn('w:type'), 'dxa')
                                    tcW.set(qn('w:w'), str(col_width))
                                    
                                # Reduce font size of all text in the cell to fit better
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        current_size_pt = run.font.size.pt if run.font.size else 11.0
                                        new_size_pt = min(9.0, current_size_pt - 2.0)
                                        run.font.size = Pt(new_size_pt)

                # Remove empty paragraphs at the beginning of the document
                while doc.paragraphs and not doc.paragraphs[0].text.strip():
                    p = doc.paragraphs[0]._element
                    p.getparent().remove(p)
                    # p._p = p._element = None (This causes AttributeError sometimes, so just remove the XML element)
                    
                doc.save(self.output_file)
            except ImportError:
                self.logger.warning("python-docx not installed, skipping page formatting")
            except Exception as format_err:
                self.logger.warning(f"Error applying page formatting: {format_err}")
            
            self.progress.emit(100)
            status = "with syntax highlighting" if self.use_highlighting else ""
            self.logger.info(f"Word conversion successful: {self.output_file}")
            self.finished.emit(True, f"Successfully converted to Word {status}:\n{self.output_file}")
            
        except Exception as e:
            self.logger.error(f"Word conversion failed: {str(e)}")
            self.finished.emit(False, f"Word conversion failed: {str(e)}")
    
    def convert_to_excel(self):
        """Convert markdown to Excel using pandas + openpyxl + BeautifulSoup + pytablewriter"""
        try:
            import markdown
            from bs4 import BeautifulSoup
            import openpyxl
            from openpyxl.styles import Font, Alignment
            from openpyxl.utils import get_column_letter
            
            try:
                import pytablewriter
                PYTABLEWRITER_AVAILABLE = True
            except ImportError:
                PYTABLEWRITER_AVAILABLE = False
                self.logger.warning("pytablewriter not available")
            
            self.logger.info(f"Starting Excel conversion: {self.input_file}")
            self.status.emit("Converting to Excel...")
            self.progress.emit(30)
            
            # Read markdown
            with open(self.input_file, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Convert to HTML
            md_content = self.fix_markdown_tables(md_content)
            html = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
            self.logger.info("Converted markdown to HTML, looking for tables")
            
            # Parse HTML and find tables, horizontal rules, headings, and text blocks
            soup = BeautifulSoup(html, 'html.parser')
            
            tracked_tags = ['table', 'hr', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'blockquote', 'pre']
            all_found = soup.find_all(tracked_tags)
            
            elements = []
            for el in all_found:
                # Check if this element is a descendant of another tracked element
                parent_is_tracked = False
                p_node = el.parent
                while p_node and p_node.name != '[document]':
                    if p_node.name in tracked_tags:
                        parent_is_tracked = True
                        break
                    p_node = p_node.parent
                
                if not parent_is_tracked:
                    elements.append(el)
            
            sections = []
            if self.excel_sheet_mode == "📊 Group by section (---)":
                current_section_elements = []
                for el in elements:
                    if el.name == 'hr':
                        if current_section_elements:
                            sections.append(current_section_elements)
                            current_section_elements = []
                    else:
                        current_section_elements.append(el)
                if current_section_elements:
                    sections.append(current_section_elements)
            elif self.excel_sheet_mode == "📊 All tables in one sheet":
                all_elements = [el for el in elements if el.name != 'hr']
                if all_elements:
                    sections.append(all_elements)
            else:
                # One table per sheet (only keep tables, ignore headings for backward compatibility)
                sections = [[t] for t in elements if t.name == 'table']
                
            table_count = sum(1 for s in sections for el in s if el.name == 'table')
            self.logger.info(f"Found {table_count} table(s) across {len(sections)} sections")
            
            if table_count > 0:
                # Process tables with formatting and merged cells
                self.status.emit(f"Found {table_count} tables, converting with formatting...")
                self.progress.emit(50)
                
                workbook = openpyxl.Workbook()
                
                anchor_map = {}
                pending_internal_links = []
                
                for section_idx, tables_in_section in enumerate(sections, 1):
                    self.status.emit(f"Processing section {section_idx}...")
                    
                    # Determine sheet name based on mode and headings
                    heading_text = ""
                    for el in tables_in_section:
                        if el.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                            text = el.get_text().strip()
                            if text:
                                heading_text = text
                                break
                                
                    if self.excel_sheet_mode == "📊 One table per sheet":
                        sheet_name = f'Table_{section_idx}' if len(sections) > 1 else 'Sheet1'
                    elif self.excel_sheet_mode == "📊 All tables in one sheet":
                        sheet_name = 'Combined_Tables'
                    else:
                        if heading_text:
                            # Clean up sheet name for Excel
                            clean_name = heading_text
                            for char in ['\\', '/', '*', '?', ':', '[', ']']:
                                clean_name = clean_name.replace(char, ' ')
                            sheet_name = clean_name.strip()[:31]
                        else:
                            sheet_name = f'Section_{section_idx}' if len(sections) > 1 else 'Sheet1'
                        
                    sheet_name = sheet_name[:31]
                    if not sheet_name:
                        sheet_name = f'Section_{section_idx}'
                    
                    if section_idx == 1:
                        worksheet = workbook.active
                        worksheet.title = sheet_name
                    else:
                        worksheet = workbook.create_sheet(title=sheet_name)
                    
                    # Track occupied cells for rowspan/colspan
                    occupied = {}
                    current_row_offset = 0
                    
                    for el in tables_in_section:
                        if el.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                            if self.excel_sheet_mode in ["📊 Group by section (---)", "📊 All tables in one sheet"]:
                                text = el.get_text().strip()
                                if text:
                                    cell = worksheet.cell(row=current_row_offset + 1, column=1, value=text)
                                    cell.font = Font(bold=True, size=12)
                                    
                                    # Track anchor ID
                                    a_tag = el.find('a', id=True)
                                    anchor_id = None
                                    if a_tag:
                                        anchor_id = a_tag.get('id')
                                    elif el.get('id'):
                                        anchor_id = el.get('id')
                                    if anchor_id:
                                        anchor_map[anchor_id] = (sheet_name, f"A{current_row_offset + 1}")
                                        
                                    current_row_offset += 2 # Leave a blank row after heading
                        elif el.name == 'table':
                            table = el
                            # Process rows
                            rows = table.find_all('tr')
                            for r_idx, tr in enumerate(rows, current_row_offset + 1):
                                c_idx = 1
                                for td in tr.find_all(['td', 'th']):
                                    # Skip cells already covered by previous rowspan
                                    while (r_idx, c_idx) in occupied:
                                        c_idx += 1
                                    
                                    # Get cell properties
                                    rowspan = int(td.get('rowspan', 1))
                                    colspan = int(td.get('colspan', 1))
                                    cell_html = str(td)
                                    cell_text = td.get_text().strip()
                                    
                                    # Write value to cell
                                    val = cell_text
                                    if val.isdigit():
                                        val = int(val)
                                    else:
                                        try:
                                            val = float(val)
                                        except ValueError:
                                            pass
                                    
                                    cell_obj = worksheet.cell(row=r_idx, column=c_idx, value=val)
                                    
                                    # Handle merged cells
                                    if rowspan > 1 or colspan > 1:
                                        last_row = r_idx + rowspan - 1
                                        last_col = c_idx + colspan - 1
                                        worksheet.merge_cells(
                                            start_row=r_idx, start_column=c_idx,
                                            end_row=last_row, end_column=last_col
                                        )
                                        
                                        # Mark all cells in merge range as occupied
                                        for r in range(r_idx, last_row + 1):
                                            for c in range(c_idx, last_col + 1):
                                                occupied[(r, c)] = True
                                    
                                    # Check for bold formatting
                                    is_bold = (
                                        '**' in cell_html or 
                                        '<strong>' in cell_html or 
                                        '<b>' in cell_html or 
                                        td.name == 'th'
                                    )
                                    
                                    # Check for italic formatting
                                    is_italic = (
                                        ('*' in cell_html and '**' not in cell_html) or
                                        '<em>' in cell_html or
                                        '<i>' in cell_html
                                    )
                                    
                                    # Apply bold/italic
                                    if is_bold and is_italic:
                                        cell_obj.font = Font(bold=True, italic=True)
                                    elif is_bold:
                                        cell_obj.font = Font(bold=True)
                                    elif is_italic:
                                        cell_obj.font = Font(italic=True)
                                        
                                    # Handle hyperlinks
                                    a_tag = td.find('a', href=True)
                                    if a_tag:
                                        href = a_tag.get('href')
                                        # Preserve existing bold/italic if any
                                        current_bold = cell_obj.font.bold
                                        current_italic = cell_obj.font.italic
                                        
                                        if href.startswith('#'):
                                            anchor_id = href[1:]
                                            pending_internal_links.append((cell_obj, anchor_id))
                                        else:
                                            cell_obj.hyperlink = href
                                        
                                        cell_obj.font = Font(color="0563C1", underline="single", bold=current_bold, italic=current_italic)
                                    
                                    # Get alignment from style or align attribute
                                    align_attr = td.get('align', '').lower()
                                    style_attr = td.get('style', '').lower()
                                    
                                    if align_attr == 'center' or 'text-align: center' in style_attr:
                                        cell_obj.alignment = Alignment(horizontal="center")
                                    elif align_attr == 'right' or 'text-align: right' in style_attr:
                                        cell_obj.alignment = Alignment(horizontal="right")
                                    elif align_attr == 'left' or 'text-align: left' in style_attr:
                                        cell_obj.alignment = Alignment(horizontal="left")
                                    
                                    # Auto-adjust column width
                                    if len(cell_text) > 0:
                                        col_letter = get_column_letter(c_idx)
                                        if col_letter in worksheet.column_dimensions:
                                            current_width = worksheet.column_dimensions[col_letter].width or 0
                                        else:
                                            current_width = 0
                                        new_width = max(current_width, len(cell_text) + 2)
                                        worksheet.column_dimensions[col_letter].width = min(new_width, 50)
                                    
                                    # Move to next column
                                    c_idx += colspan
                                    
                            if self.excel_sheet_mode in ["📊 Group by section (---)", "📊 All tables in one sheet"]:
                                current_row_offset += len(rows) + 3 # Separate tables by 3 lines
                        else:
                            # It's a text block (p, ul, ol, blockquote, pre)
                            if self.excel_sheet_mode in ["📊 Group by section (---)", "📊 All tables in one sheet"]:
                                text = el.get_text()
                                lines = [line.strip() for line in text.split('\n') if line.strip()]
                                if lines:
                                    for line in lines:
                                        if el.name in ['ul', 'ol']:
                                            line = "• " + line
                                        elif el.name == 'blockquote':
                                            line = "> " + line
                                        cell = worksheet.cell(row=current_row_offset + 1, column=1, value=line)
                                        
                                        # Track anchor ID in text block
                                        a_tag = el.find('a', id=True)
                                        anchor_id = None
                                        if a_tag:
                                            anchor_id = a_tag.get('id')
                                        elif el.get('id'):
                                            anchor_id = el.get('id')
                                        if anchor_id:
                                            anchor_map[anchor_id] = (sheet_name, f"A{current_row_offset + 1}")
                                        
                                        current_row_offset += 1
                                    current_row_offset += 1 # Add blank line after text block
                
                # Resolve internal links
                for cell_obj, anchor_id in pending_internal_links:
                    if anchor_id in anchor_map:
                        target_sheet, target_coord = anchor_map[anchor_id]
                        cell_obj.hyperlink = f"#'{target_sheet}'!{target_coord}"
                        
                workbook.save(self.output_file)
                
                self.progress.emit(100)
                self.logger.info(f"Excel conversion successful: {table_count} tables to {self.output_file}")
                self.finished.emit(True, f"Successfully converted {table_count} table(s) to Excel:\n{self.output_file}")
                return
            
            # No tables found, try pytablewriter for structured data
            if PYTABLEWRITER_AVAILABLE:
                self.status.emit("No tables found, trying pytablewriter for structured data...")
                self.logger.info("No tables found, trying pytablewriter")
                
                # Parse markdown for lists and key-value pairs
                lines = md_content.split('\n')
                data = []
                for line in lines[:200]:
                    if ':' in line and not line.startswith('#'):
                        parts = line.split(':', 1)
                        key = parts[0].strip()
                        # Strict check: key must not be too long, and should not be a normal sentence with spaces
                        if len(key) < 40 and not ' ' in key.strip('-* '):
                            data.append([key.strip('-* '), parts[1].strip()])
                    elif line.startswith('- ') or line.startswith('* '):
                        data.append(['List Item', line[2:].strip()])
                
                if data:
                    writer = pytablewriter.ExcelXlsxTableWriter()
                    writer.open(self.output_file)
                    writer.headers = ['Key', 'Value']
                    writer.value_matrix = data
                    writer.write_table()
                    writer.close()
                    
                    self.progress.emit(100)
                    self.logger.info(f"Pytablewriter conversion successful: {self.output_file}")
                    self.finished.emit(True, f"Converted structured data to Excel using pytablewriter:\n{self.output_file}")
                    return
            
            self.logger.warning("No tables or structured data found")
            self.finished.emit(False, "No tables or structured data found in markdown file")
            
        except Exception as e:
            self.logger.error(f"Excel conversion failed: {str(e)}")
            self.finished.emit(False, f"Excel conversion failed: {str(e)}")

    def convert_docx_to_md(self):
        """Convert Word document to Markdown using pypandoc"""
        try:
            import pypandoc
            
            self.logger.info(f"Starting reverse Word conversion: {self.input_file}")
            self.status.emit("Converting Word to Markdown...")
            self.progress.emit(30)
            
            extra_args = []
            if hasattr(self, 'extract_images') and self.extract_images:
                # Pandoc can extract images to a directory
                media_dir = Path(self.output_file).parent / "images"
                extra_args.append(f'--extract-media={media_dir}')
            
            pypandoc.convert_file(
                self.input_file,
                'gfm',
                format='docx',
                outputfile=self.output_file,
                extra_args=extra_args
            )
            
            self.progress.emit(100)
            self.logger.info(f"Reverse Word conversion successful: {self.output_file}")
            self.finished.emit(True, f"Successfully converted Word to Markdown:\n{self.output_file}")
            
        except Exception as e:
            self.logger.error(f"Reverse Word conversion failed: {str(e)}")
            self.finished.emit(False, f"Reverse Word conversion failed: {str(e)}")

    def convert_xlsx_to_md(self):
        """Convert Excel to Markdown preserving sheets, headings, and formatting"""
        try:
            import openpyxl
            
            self.logger.info(f"Starting reverse Excel conversion: {self.input_file}")
            self.status.emit("Loading Excel file...")
            self.progress.emit(20)
            
            wb = openpyxl.load_workbook(self.input_file, data_only=True)
            markdown_content = []
            
            sheet_count = len(wb.sheetnames)
            for sheet_idx, sheet_name in enumerate(wb.sheetnames):
                self.status.emit(f"Parsing sheet {sheet_idx+1}/{sheet_count}: {sheet_name}")
                self.progress.emit(20 + int(70 * (sheet_idx / sheet_count)))
                
                sheet = wb[sheet_name]
                if sheet_idx > 0:
                    markdown_content.append("\n---\n")
                
                markdown_content.append(f"## Sheet: {sheet_name}\n")
                
                # Analyze rows to determine blocks (tables vs text)
                current_block = []
                for row in sheet.iter_rows():
                    # Count non-empty cells
                    filled_cells = []
                    for cell in row:
                        if cell.value is not None and str(cell.value).strip() != "":
                            filled_cells.append(cell)
                            
                    if len(filled_cells) == 0:
                        # Empty row breaks current block
                        if current_block:
                            self._flush_excel_block(current_block, markdown_content)
                            current_block = []
                        markdown_content.append("")
                    elif len(filled_cells) == 1:
                        # Single cell might be a heading or just text
                        if current_block:
                            self._flush_excel_block(current_block, markdown_content)
                            current_block = []
                        
                        cell = filled_cells[0]
                        val = str(cell.value).strip()
                        
                        # Formatting
                        prefix = ""
                        suffix = ""
                        if cell.font:
                            if cell.font.b or (cell.font.sz and cell.font.sz > 12):
                                prefix = "### "
                                suffix = ""
                            elif cell.font.i:
                                prefix = "*"
                                suffix = "*"
                                
                        if cell.hyperlink:
                            val = f"[{val}]({cell.hyperlink.target})"
                            
                        markdown_content.append(f"{prefix}{val}{suffix}\n")
                    else:
                        # Multiple cells, part of a table
                        current_block.append(row)
                
                # Flush remaining block
                if current_block:
                    self._flush_excel_block(current_block, markdown_content)
                    
            with open(self.output_file, 'w', encoding='utf-8') as f:
                f.write('\n'.join(markdown_content))
                
            self.progress.emit(100)
            self.logger.info(f"Reverse Excel conversion successful: {self.output_file}")
            self.finished.emit(True, f"Successfully converted Excel to Markdown:\n{self.output_file}")
            
        except Exception as e:
            self.logger.error(f"Reverse Excel conversion failed: {str(e)}")
            self.finished.emit(False, f"Reverse Excel conversion failed: {str(e)}")

    def _flush_excel_block(self, rows, markdown_content):
        if not rows: return
        
        # Determine columns length
        max_col = max(len(r) for r in rows)
        
        # Build table
        table_lines = []
        for row_idx, row in enumerate(rows):
            line_parts = []
            for col_idx in range(max_col):
                if col_idx < len(row):
                    cell = row[col_idx]
                    val = str(cell.value).strip() if cell.value is not None else ""
                    
                    if val != "" and cell.font:
                        if cell.font.b: val = f"**{val}**"
                        elif cell.font.i: val = f"*{val}*"
                    if val != "" and cell.hyperlink:
                        val = f"[{val}]({cell.hyperlink.target})"
                        
                    # Escape pipes
                    val = val.replace('|', '\\|')
                    # Remove newlines inside cells
                    val = val.replace('\n', ' <br> ')
                    line_parts.append(val)
                else:
                    line_parts.append("")
            
            table_lines.append("| " + " | ".join(line_parts) + " |")
            
            # Add separator after header row
            if row_idx == 0:
                table_lines.append("|" + "|".join(["---" for _ in range(max_col)]) + "|")
                
        markdown_content.extend(table_lines)
        markdown_content.append("\n")